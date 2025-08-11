from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import re
import os

# Document type keywords (unchanged)
DOC_TYPE_KEYWORDS = {
    "Articles of Association": [
        "articles of association", "aoa", "article", "share capital", "object clause"
    ],
    "Memorandum of Association": [
        "memorandum of association", "moa", "memorandum", "subscribers"
    ],
    "Incorporation Application Form": [
        "incorporation application", "application for incorporation", "registration application"
    ],
    "UBO Declaration": [
        "ultimate beneficial owner", "ubo declaration", "beneficial owner"
    ],
    "Board Resolution": [
        "board resolution", "resolved that", "by the board", "board meeting"
    ],
    "Register of Members and Directors": [
        "register of members", "register of directors", "register of members and directors"
    ],
    "Shareholder Resolution": [
        "shareholder resolution", "resolution of the shareholders"
    ],
    "Employment Contract": [
        "employment contract", "employee", "employer", "job title", "probation"
    ],
    "Data Protection Policy": [
        "data protection", "privacy policy", "appropriate policy document", "dpr"
    ],
    "License Application Form": [
        "license application", "licensing application", "business license"
    ],
    "Business Plan": [
        "business plan", "strategic plan", "business strategy"
    ],
    "Financial Projections": [
        "financial projections", "financial forecast", "financial plan"
    ],
    "Change of Registered Address Notice": [
        "change of registered address", "address change notice", "registered address change", "address notification"
    ],
    "Unknown": []
}

def extract_paragraphs_from_docx(path):
    doc = Document(path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = "\n".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                paragraphs.append(" | ".join(row_texts))
    return paragraphs

def chunk_paragraphs(paragraphs, max_chars=2000):
    chunks = []
    current = []
    current_len = 0
    for p in paragraphs:
        p_len = len(p)
        if current_len + p_len + 1 > max_chars:
            if current:
                chunks.append("\n\n".join(current))
            current = [p]
            current_len = p_len
        else:
            current.append(p)
            current_len += p_len + 1
    if current:
        chunks.append("\n\n".join(current))
    return chunks

def detect_doc_type(text, threshold=1):
    text_low = text.lower()
    scores = {}
    for doc_type, kws in DOC_TYPE_KEYWORDS.items():
        score = 0
        for kw in kws:
            if kw.lower() in text_low:
                score += 1
        scores[doc_type] = score
    best = max(scores.items(), key=lambda x: x[1])
    if best[1] >= threshold and best[0] != "Unknown":
        return best[0], scores
    return "Unknown", scores

def detect_red_flags(text, doc_type, rag_system=None):
    issues = []
    text_lower = text.lower()
    
    # Rule-based checks with ADGM law citations and alternative wording
    jurisdiction_patterns = [
        r"UAE Federal Courts",
        r"Federal Court of UAE", 
        r"Abu Dhabi Court",
        r"Dubai Courts"
    ]
    for pattern in jurisdiction_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            issues.append({
                "type": "jurisdiction_issue",
                "severity": "High",
                "issue": f"References to {match.group()} instead of ADGM Courts",
                "suggestion": "Replace with ADGM Courts reference",
                "section": match.group(),
                "adgm_source": "ADGM Companies Regulations 2020, Art. 6",
                "alternative": "The courts of the Abu Dhabi Global Market shall have exclusive jurisdiction to settle any dispute arising out of or in connection with this document."
            })
    
    # Define required clauses with specific ADGM-compliant suggestions
    required_clauses = {
        "governing law": {
            "suggestion": "Add governing law clause: 'This document shall be governed by and construed in accordance with the laws of the Abu Dhabi Global Market.'",
            "reference": "ADGM Companies Regulations 2020, Art. 6",
            "alternative": "This Agreement shall be governed by and construed in accordance with the laws of Abu Dhabi Global Market (ADGM)."
        },
        "jurisdiction": {
            "suggestion": "Add jurisdiction clause: 'The courts of the Abu Dhabi Global Market shall have exclusive jurisdiction to settle any dispute arising out of or in connection with this document.'",
            "reference": "ADGM Companies Regulations 2020, Art. 6",
            "alternative": "Any dispute arising out of or in connection with this Agreement, including any question regarding its existence, validity or termination, shall be subject to the exclusive jurisdiction of the courts of Abu Dhabi Global Market."
        },
        "dispute resolution": {
            "suggestion": "Add dispute resolution clause: 'Any dispute arising out of or in connection with this document shall be resolved through arbitration in accordance with ADGM Arbitration Regulations.'",
            "reference": "ADGM Arbitration Regulations 2015",
            "alternative": "All disputes arising out of or in connection with this Agreement shall be finally settled under the Rules of Arbitration of the ADGM Arbitration Centre by one or more arbitrators appointed in accordance with the said Rules."
        }
    }
    
    for clause, details in required_clauses.items():
        if clause not in text_lower:
            issues.append({
                "type": "missing_clause",
                "severity": "Medium",
                "issue": f"Missing {clause} clause",
                "suggestion": f"{details['suggestion']}\n\nAlternative wording: {details['alternative']}",
                "section": "N/A",
                "adgm_source": details['reference']
            })
    
    # Check for ambiguous language with ADGM citation and alternatives
    ambiguous_patterns = [
        (r"may be", "shall be", "is"),
        (r"could be", "shall be", "is"),
        (r"might", "shall", "will"),
        (r"possibly", "definitely", "certainly")
    ]
    for pattern, definitive, alternative in ambiguous_patterns:
        matches = re.finditer(pattern, text_lower)
        for match in matches:
            ambiguous_text = match.group()
            # Get surrounding context (up to 10 words before and after)
            start = max(0, match.start() - 50)
            end = min(len(text_lower), match.end() + 50)
            context = text_lower[start:end]
            
            issues.append({
                "type": "ambiguous_language",
                "severity": "Low",
                "issue": f"Ambiguous language: '{ambiguous_text}'",
                "suggestion": f"Replace with definitive language such as '{definitive}' for legal certainty",
                "section": context,
                "adgm_source": "ADGM Legal Documentation Guidelines",
                "alternative": f"Consider using '{definitive}' or '{alternative}' instead of '{ambiguous_text}'"
            })
    
    # Check for missing signatory sections with ADGM citation and alternative wording
    signatory_patterns = [
        r"signature",
        r"signatory", 
        r"signed by",
        r"authorized signatory",
        r"witness"
    ]
    has_signatory = any(re.search(pattern, text_lower) for pattern in signatory_patterns)
    if not has_signatory and doc_type in ["Articles of Association", "Memorandum of Association", "Board Resolution", "Shareholder Resolution"]:
        adgm_source = "ADGM Companies Regulations 2020"
        alternative = ""
        
        if doc_type == "Articles of Association" or doc_type == "Memorandum of Association":
            adgm_source = "ADGM Companies Regulations 2020, Art. 12-13"
            alternative = "IN WITNESS WHEREOF, the Subscriber has executed these Articles of Association on the date first written above.\n\nName: [Full Name]\nPosition: [Director/Authorized Signatory]\nSignature: _________________\n\nWitness:\nName: [Full Name]\nSignature: _________________"
        elif doc_type == "Board Resolution" or doc_type == "Shareholder Resolution":
            adgm_source = "ADGM Companies Regulations 2020, Art. 154-158"
            alternative = "CERTIFIED TRUE COPY\n\nChairperson of the Meeting:\nName: [Full Name]\nSignature: _________________\nDate: [DD/MM/YYYY]\n\nDirector/Company Secretary:\nName: [Full Name]\nSignature: _________________\nDate: [DD/MM/YYYY]"
            
        issues.append({
            "type": "missing_signatory",
            "severity": "High",
            "issue": "Missing signatory section or signature block",
            "suggestion": "Add proper signature section with authorized signatories as per ADGM requirements",
            "section": "N/A",
            "adgm_source": adgm_source,
            "alternative": alternative
        })
    
    # Check for template compliance issues with ADGM citation
    template_indicators = [
        r"insert.*here",
        r"\[.*\]",
        r"\{.*\}",
        r"placeholder",
        r"template"
    ]
    for pattern in template_indicators:
        matches = re.finditer(pattern, text_lower)
        for match in matches:
            placeholder_text = match.group()
            # Generate appropriate alternative text based on the placeholder
            alternative = ""
            
            # Try to determine what kind of content is expected
            if "name" in placeholder_text.lower():
                alternative = "[Full legal name as registered with ADGM]"
            elif "date" in placeholder_text.lower():
                alternative = "[DD/MM/YYYY format as per ADGM requirements]"
            elif "address" in placeholder_text.lower():
                alternative = "[Complete registered address including ADGM building/floor/unit]"
            elif "amount" in placeholder_text.lower() or "sum" in placeholder_text.lower():
                alternative = "[Exact amount in numbers and words]"
            elif "signature" in placeholder_text.lower():
                alternative = "[Authorized signature with name and title]"
            else:
                alternative = "[Appropriate content as per ADGM requirements]"
                
            issues.append({
                "type": "template_placeholder",
                "severity": "Medium",
                "issue": f"Template placeholder found: {placeholder_text}",
                "suggestion": "Replace template placeholders with actual content before submission",
                "section": placeholder_text,
                "adgm_source": "ADGM Documentation Requirements",
                "alternative": alternative
            })
    
    # Check for formatting issues
    # More selective formatting issues with ADGM citations and alternatives
    formatting_issues = [
        # Only detect multiple consecutive empty paragraphs (3+) to reduce noise
        (r"(^\s*$\n){3,}", "Multiple consecutive empty paragraphs", "ADGM Documentation Style Guide", "Consider using section breaks or proper spacing between paragraphs"),
        
        # Only detect ALL CAPS text that's likely a heading or emphasis (10+ characters)
        (r"[A-Z]{10,}", "Excessive capitalization in text", "ADGM Documentation Style Guide", "Use proper heading styles or bold formatting instead of all caps"),
        
        # Date format check with ADGM citation
        (r"\d{4}-\d{2}-\d{2}", "Date format should be DD/MM/YYYY for ADGM documents", "ADGM Documentation Requirements", "Format as DD/MM/YYYY (e.g., 01/01/2023 instead of 2023-01-01)")
    ]
    for pattern, description, adgm_source, alternative in formatting_issues:
        matches = re.finditer(pattern, text, re.MULTILINE)
        for match in matches:
            issues.append({
                "type": "formatting_issue",
                "severity": "Low",
                "issue": description,
                "suggestion": "Review and correct document formatting according to ADGM standards",
                "section": match.group(),
                "adgm_source": adgm_source,
                "alternative": alternative
            })
    
    # Enhance with RAG if available - include ADGM sources and alternatives
    if rag_system and doc_type != "Unknown":
        rag_response = rag_system.analyze_document(text, doc_type)
        if 'flags' in rag_response:
            for flag in rag_response['flags']:
                # Extract ADGM source if available in the RAG response
                adgm_source = flag.get('adgm_source', '')
                if not adgm_source and 'source' in flag:
                    adgm_source = flag.get('source', '')
                if not adgm_source:
                    adgm_source = "ADGM Documentation Guidelines"
                
                # Extract or generate alternative wording
                alternative = flag.get('alternative', '')
                if not alternative and flag.get('suggestion', ''):
                    # Generate a simple alternative based on the suggestion
                    suggestion = flag.get('suggestion', '')
                    if len(suggestion) > 10:  # Only if suggestion is substantial
                        alternative = suggestion
                
                issues.append({
                    "type": flag.get('type', 'rag_issue'),
                    "severity": flag.get('severity', 'Medium'),
                    "issue": flag.get('reason', 'Unknown issue'),
                    "suggestion": flag.get('suggestion', ''),
                    "section": flag.get('snippet', 'N/A'),
                    "adgm_source": adgm_source,
                    "alternative": alternative
                })
    
    return issues

def add_comments_to_docx(filepath, issues):
    """
    Add inline comments to .docx file for each issue.
    Optimized to reduce repetitive comments and ensure proper positioning.
    """
    try:
        doc = Document(filepath)
        modified = False
        
        # Filter out repetitive formatting issues at the beginning
        filtered_issues = []
        for issue in issues:
            issue_type = issue.get('type', 'unknown')
            # Skip formatting issues that are too repetitive
            if issue_type == 'formatting_issue' and any(x in issue.get('issue', '').lower() for x in ['whitespace', 'spacing', 'capitalization']):
                continue
            filtered_issues.append(issue)
        
        # If no issues remain after filtering, add a summary comment to the first paragraph
        if not filtered_issues:
            error_comment = "✅ No compliance issues detected. Document reviewed by ADGM Corporate Agent."
            if doc.paragraphs:
                add_comment_to_paragraph(doc.paragraphs[0], error_comment)
            output_path = filepath.replace(".docx", "_reviewed.docx")
            doc.save(output_path)
            return output_path
        
        # Check if original issues contained an error from LLM/RAG
        error_found = False
        for issue in issues:
            if issue.get('type') == 'error' or 'error' in issue.get('issue', '').lower():
                error_comment = f"❌ AI/RAG analysis failed: {issue.get('issue', 'Unknown error')}. Using rule-based checks only."
                if doc.paragraphs:
                    add_comment_to_paragraph(doc.paragraphs[0], error_comment)
                error_found = True
                modified = True
                break
            
        # Group issues by section to avoid multiple comments on the same text
        issue_groups = {}
        for issue in filtered_issues:
            section = issue.get('section', 'N/A')
            if section not in issue_groups:
                issue_groups[section] = []
            issue_groups[section].append(issue)
        
        # Process each section's issues
        for section, section_issues in issue_groups.items():
            # Handle issues without a specific section
            if section == 'N/A':
                # For issues without a specific section, add to the first paragraph
                if doc.paragraphs:
                    comment_parts = []
                    for issue in section_issues:
                        # Include ADGM law citation
                        adgm_source = issue.get('adgm_source', '')
                        reference = f" (Per {adgm_source})" if adgm_source else ""
                        
                        # Include alternative wording if available
                        alternative = issue.get('alternative', '')
                        alt_text = f"\n🔄 ALTERNATIVE WORDING: {alternative}" if alternative else ""
                        
                        comment_parts.append(f"🚨 ISSUE: {issue['issue']}{reference}\n💡 SUGGESTION: {issue['suggestion']}{alt_text}\n📋 SEVERITY: {issue['severity']}")
                    
                    if comment_parts:
                        comment_text = "\n\n".join(comment_parts)
                        add_comment_to_paragraph(doc.paragraphs[0], comment_text)
                        modified = True
                continue
                
            # Deduplicate issues with similar suggestions
            unique_issues = []
            suggestion_set = set()
            for issue in section_issues:
                suggestion = issue.get('suggestion', '')
                if suggestion not in suggestion_set:
                    suggestion_set.add(suggestion)
                    unique_issues.append(issue)
            
            # Skip if no unique issues after deduplication
            if not unique_issues:
                continue
                
            comment_added = False
            
            # Try to find exact match for the section text
            for para in doc.paragraphs:
                if section != 'N/A' and section.lower() in para.text.lower():
                    # Combine all issues for this section into one comment
                    comment_parts = []
                    for issue in unique_issues:
                        # Include ADGM law citation
                        adgm_source = issue.get('adgm_source', '')
                        reference = f" (Per {adgm_source})" if adgm_source else ""
                        
                        # Include alternative wording if available
                        alternative = issue.get('alternative', '')
                        alt_text = f"\n🔄 ALTERNATIVE WORDING: {alternative}" if alternative else ""
                        
                        comment_parts.append(f"🚨 ISSUE: {issue['issue']}{reference}\n💡 SUGGESTION: {issue['suggestion']}{alt_text}\n📋 SEVERITY: {issue['severity']}")
                    
                    comment_text = "\n\n".join(comment_parts)
                    add_comment_to_paragraph(para, comment_text)
                    comment_added = True
                    modified = True
                    break
            
            # If exact match not found, try to find a paragraph with similar content
            if not comment_added and doc.paragraphs and section != 'N/A':
                best_match = None
                highest_similarity = 0
                
                # Find paragraph with highest similarity to the section
                for para in doc.paragraphs:
                    if para.text.strip():
                        # Simple similarity measure - count of common words
                        section_words = set(section.lower().split())
                        para_words = set(para.text.lower().split())
                        common_words = section_words.intersection(para_words)
                        
                        # Calculate similarity score
                        if len(section_words) > 0:
                            similarity = len(common_words) / len(section_words)
                            
                            # Prioritize paragraphs that contain key words from the section
                            if similarity > highest_similarity and similarity > 0.4:  # 40% threshold
                                highest_similarity = similarity
                                best_match = para
                
                if best_match:
                    # Combine all issues for this section into one comment
                    comment_parts = []
                    for issue in unique_issues:
                        # Include ADGM law citation
                        adgm_source = issue.get('adgm_source', '')
                        reference = f" (Per {adgm_source})" if adgm_source else ""
                        
                        # Include alternative wording if available
                        alternative = issue.get('alternative', '')
                        alt_text = f"\n🔄 ALTERNATIVE WORDING: {alternative}" if alternative else ""
                        
                        comment_parts.append(f"🚨 ISSUE: {issue['issue']}{reference}\n💡 SUGGESTION: {issue['suggestion']}{alt_text}\n📋 SEVERITY: {issue['severity']}")
                    
                    comment_text = "\n\n".join(comment_parts)
                    add_comment_to_paragraph(best_match, comment_text)
                    comment_added = True
                    modified = True
        
        # Always save a reviewed document, even if no comments were added
        output_path = filepath.replace(".docx", "_reviewed.docx")
        doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"Error processing document {filepath}: {e}")
        return filepath

def add_comment_to_paragraph(paragraph, comment_text):
    """
    Add a comment to a specific paragraph using python-docx.
    Uses the built-in comment functionality in python-docx 1.2.0.
    """
    try:
        # Get the document from the paragraph
        document = paragraph.part.document
        
        # Use the built-in add_comment method in python-docx 1.2.0
        # We need to pass the paragraph's runs to the add_comment method
        if not paragraph.runs:
            # If paragraph has no runs, add a run to make sure we can attach a comment
            paragraph.add_run()
        
        # Add the comment to the document using the built-in method
        document.add_comment(
            runs=paragraph.runs,
            text=comment_text,
            author="ADGM Corporate Agent",
            initials="ADGM"
        )
    except Exception as e:
        print(f"Warning: Could not add comment to document: {e}")
        return

def build_file_summary(filepath, rag_system=None, max_preview_chars=1200):
    try:
        paragraphs = extract_paragraphs_from_docx(filepath)
        full_text = "\n\n".join(paragraphs)
        preview = full_text[:max_preview_chars] + ("..." if len(full_text) > max_preview_chars else "")
        chunks = chunk_paragraphs(paragraphs, max_chars=2000)
        doc_type, scores = detect_doc_type(full_text)
        issues = detect_red_flags(full_text, doc_type, rag_system)
        
        return {
            "filename": filepath,
            "detected_type": doc_type,
            "paragraph_count": len(paragraphs),
            "chunk_count": len(chunks),
            "preview": preview,
            "full_text": full_text,
            "keyword_scores": scores,
            "issues": issues
        }
    except Exception as e:
        return {
            "filename": filepath,
            "detected_type": "Error",
            "paragraph_count": 0,
            "chunk_count": 0,
            "preview": f"Error processing file: {str(e)}",
            "full_text": "",
            "keyword_scores": {},
            "issues": [{"type": "error", "severity": "High", "issue": f"File processing failed: {str(e)}", "suggestion": "Check file format", "section": "N/A"}]
        }

def save_json_summary(summary_list, out_path):
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(summary_list, f, indent=2, ensure_ascii=False)